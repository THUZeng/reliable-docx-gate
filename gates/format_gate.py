# -*- coding: utf-8 -*-
"""
标书格式两道闸门 (format_gate) —— 交付前必跑，两道全绿才算「格式对」。
- Gate A · 结构闸门 (读 XML 硬事实, 可靠层)：目录/封面/换页/标题样式/编号连续/孤儿样式/断句/偏离表全满足/字体。
- Gate B · 视觉闸门 (渲染成 PNG, 由 Claude 本人当 VLM 亲眼看, 不外包千问)：豆腐块/表格出界/标题孤行/图压字/页码连续。

清单每条都「有据可循」，来源见 SKILL.md「两道格式闸门」章的引用。
XML 只对「声明层」可靠, 对「渲染层」(缺字豆腐块/是否出界/分页)不可靠 —— 所以 Gate B 必跑, 不能只靠 Gate A。

用法:
    from format_gate import run_gate
    run_gate("商务技术文件.docx", reference="磋商文件.docx", out_png_dir="/tmp/gateB")
    # -> 打印 Gate A 报告 + 生成 Gate B 的 PNG 路径列表(Claude 逐张 Read 去看)
"""
import os, re, subprocess, glob
from collections import Counter
from docx import Document
from docx.oxml.ns import qn

TERMINAL = "。！？；：）」》】.!?;:)"          # 段末正常应有的收尾标点
GOOD_STYLES = {"Normal", "Normal Indent", "标题", "正文", "正文2", "CM14",
               "Heading 1", "Heading 2", "Heading 3", "toc 1", "toc 2", "TOC Heading",
               "List Paragraph", "List Paragraph1", "部分1", "Plain Text"}
SEC_RE = re.compile(r"^\s*([一二三四五六七八九十]+)\s*、")     # 一、二、… 中文顺序号

def _xml(doc):
    return doc.element.body.xml

def _has_toc_field(doc):
    x = _xml(doc)
    return ("instrText" in x and "TOC" in x) or ("fldSimple" in x and "TOC" in x)

def _has_manual_toc(doc):
    # 有「目录」标题, 且随后若干段以数字页码结尾(制表位+页码) —— 手工目录也算有目录
    ps = doc.paragraphs
    for i, p in enumerate(ps):
        if p.text.strip() in ("目录", "目 录"):
            tail = " ".join(q.text for q in ps[i+1:i+15])
            if re.search(r"\d+\s*$", tail) or "\t" in tail:
                return True
    return False

def _effective_eastasia_font(p):
    for r in p.runs:
        rpr = r._element.rPr
        if rpr is not None:
            f = rpr.find(qn('w:rFonts'))
            if f is not None and f.get(qn('w:eastAsia')):
                return f.get(qn('w:eastAsia'))
    return None  # 继承样式, 未显式声明

def gate_a(docx_path, reference=None):
    d = Document(docx_path)
    ps = d.paragraphs
    x = _xml(d)
    R = []  # (ok, name, detail)  ok: True/False/None(=需人工/Gate B)

    # 1. 换页 (页码不连续/内容挤一坨的根因)  —— 来源: CSDN废标雷区, 100条规范
    #    两种都算真换页: 显式断行 <w:br type=page> 和 段前分页 <w:pageBreakBefore>
    pb = x.count('w:type="page"') + x.count('<w:pageBreakBefore')
    R.append((pb > 0, "有换页(page break)", f"{pb} 处"))

    # 2. 目录 (自动生成优先; 手工目录也可)  —— 来源: 知乎排版, CSDN
    toc = _has_toc_field(d) or _has_manual_toc(d)
    R.append((toc, "有目录", "TOC域" if _has_toc_field(d) else ("手工目录" if _has_manual_toc(d) else "无")))

    # 3. 封面 (首段大字居中标题, body 之前)  —— 来源: 知乎/排版规范
    cover = False
    for p in ps[:6]:
        sz = p.runs[0].font.size.pt if (p.runs and p.runs[0].font.size) else 0
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        if p.text.strip() and sz and sz >= 18 and p.alignment == WD_ALIGN_PARAGRAPH.CENTER:
            cover = True; break
    R.append((cover, "有封面(首页大字居中标题)", "是" if cover else "未检出"))

    # 4. 标题样式 (全 Normal = 断结构层)  —— 常见坑: 读 .text 重建会把样式压平成全 Normal
    sc = Counter(p.style.name for p in ps)
    head_styles = sum(v for k, v in sc.items() if ("eading" in k) or k.startswith("标题"))
    ratio_normal = sc.get("Normal", 0) / max(1, len(ps))
    R.append((head_styles > 0, "用了真标题样式(非全Normal)",
              f"标题样式段={head_styles}, Normal占比={ratio_normal:.0%}"))

    # 5. 孤儿/外来样式 (粘贴带入, 如 纯文本_0_0 / AONormal)
    orphans = [k for k in sc if k not in GOOD_STYLES and sc[k] <= 3 and not k.startswith("Heading")]
    R.append((len(orphans) == 0, "无孤儿粘贴样式", ("干净" if not orphans else "可疑: " + ", ".join(orphans))))

    # 6. 编号连续 (缺 三、/ 重复 5、)  —— 来源: CSDN 100条 第21条
    #    有标题样式时只查「章节标题」的编号(排除廉洁承诺书等嵌套一二三小列表), 避免误报
    seq = []
    zh = "一二三四五六七八九十"
    head_ps = [p for p in ps if ("eading" in p.style.name) or p.style.name.startswith("标题")]
    target = head_ps if head_ps else ps
    for p in target:
        m = SEC_RE.match(p.text)
        if m and len(p.text.strip()) < 45:
            seq.append(zh.index(m.group(1)[0]) + 1 if m.group(1) in zh else None)
    # 找主序列(投标函一二三…)的最长连续段, 报缺号/重号
    dup = [n for n in set(seq) if seq.count(n) > 1]
    gap = []
    s = [n for n in seq if n]
    if s:
        for n in range(min(s), max(s) + 1):
            if n not in s: gap.append(n)
    ok6 = (not dup and not gap)
    R.append((ok6, "章节中文编号连续无缺无重",
              f"序列={seq[:12]}{' 重号'+str(dup) if dup else ''}{' 缺号'+str(gap) if gap else ''}"))

    # 7. 跨行断句 (段末无收尾标点 + 很短 => OCR/文本重建把一行拆成一段)
    #    跳过: 封面区/标题样式/带冒号的标签行/日期/图表题注/TOC占位, 免误报
    DATE_RE = re.compile(r"\d+\s*年|\d+\s*月")
    splits = []
    for i, p in enumerate(ps):
        t = p.text.strip()
        if i < 18: continue
        if ("eading" in p.style.name) or p.style.name.startswith("标题"): continue
        if ("：" in t) or (":" in t) or DATE_RE.search(t): continue
        if t.startswith("图") or t.startswith("表") or ("更新域" in t) or ("目录" in t): continue
        if 0 < len(t) <= 28 and t[-1] not in TERMINAL:
            nxt = ps[i+1].text.strip() if i+1 < len(ps) else ""
            if nxt and nxt[0] not in "一二三四五六七八九十（(0123456789注图表":
                splits.append(t[-14:])
    R.append((len(splits) == 0, "无跨行断句(段末标点正常)",
              ("正常" if not splits else f"{len(splits)}处可疑, 例: " + " / ".join(splits[:3]))))

    # 8. 偏离表全「满足」(▲负偏离=废标)  —— 来源: 知乎废标汇总
    dev_ok, dev_detail = None, "未见偏离表"
    for tb in d.tables:
        header = " ".join(c.text for c in tb.rows[0].cells)
        if "偏离" in header:
            col = next((j for j, c in enumerate(tb.rows[0].cells) if "偏离" in c.text), None)
            if col is not None:
                vals = [tb.rows[i].cells[col].text.strip() for i in range(1, len(tb.rows))]
                POS = ("满足", "无偏离", "正偏离", "响应")
                # ⚠「不满足」含子串「满足」——必须先查否定词, 否则负偏离(废标)会被漏(变异测试查出)
                NEG = ("不满足", "未满足", "负偏离", "不响应", "部分满足", "有偏离", "不符", "偏高", "偏低")
                bad = [v for v in vals if v and (any(n in v for n in NEG) or not any(p in v for p in POS))]
                dev_ok = (len(bad) == 0)
                dev_detail = "全满足" if dev_ok else f"疑负偏离: {bad[:3]}"
            break
    R.append((dev_ok, "偏离表全满足(无负偏离)", dev_detail))

    # 8b. 无残留占位符 (未填的 XX/XXX/见投标文件第页/空括号/下划线) —— 常见坑: 只往空单元格填, 会漏掉占位文本(XX/见…第页)
    #     白名单: 「（此处加盖电子公章）」是电子章标注(留)、「〔2017〕/【项目编号:...】」是真实公文号(留)
    _ph_pats = [
        (re.compile(r"X{2,}"), "XX/XXX"),
        (re.compile(r"见投标文件第\s*页"), "见投标文件第页"),
        (re.compile(r"〔\s*〕"), "空〔〕"),
        (re.compile(r"【\s*】"), "空【】"),
        (re.compile(r"[_＿]{4,}"), "____下划线"),
    ]
    def _scan_placeholder(text):
        out = []
        for rgx, lbl in _ph_pats:
            if rgx.search(text):
                out.append(lbl)
        for m in re.finditer(r"（此处[^）]*）", text):   # （此处…） 除电子章标注外均为占位
            if "加盖电子公章" not in m.group(0):
                out.append("（此处…）")
        return out
    all_texts = [p.text for p in ps]
    for tb in d.tables:
        for row in tb.rows:
            for c in row.cells:
                all_texts.append(c.text)
    ph_hits = []
    for text in all_texts:
        for lbl in _scan_placeholder(text):
            ph_hits.append((lbl, text.strip().replace("\n", " ")[:24]))
    R.append((len(ph_hits) == 0, "无残留占位符(XX/见投标文件第页/空括号)",
              ("干净" if not ph_hits else f"{len(ph_hits)}处: " + " / ".join(f"[{l}]{s}" for l, s in ph_hits[:4]))))

    # 9. 正文中文字体 (排版规范: 正文小四宋体; 声明层, 渲染字体见 Gate B)
    fonts = Counter()
    for p in ps:
        f = _effective_eastasia_font(p)
        if f and len(p.text.strip()) > 10: fonts[f] += 1
    R.append((None, "正文中文字体(声明层)", dict(fonts.most_common(4)) or "多为继承样式, 需Gate B看渲染"))

    # 与范文/磋商指纹对比
    if reference and os.path.exists(reference):
        rd = Document(reference); rx = _xml(rd)
        ref_pb = rx.count('w:type="page"'); ref_toc = _has_toc_field(rd) or _has_manual_toc(rd)
        rsc = Counter(p.style.name for p in rd.paragraphs)
        ref_head = sum(v for k, v in rsc.items() if "eading" in k or k.startswith("标题"))
        diffs = []
        if ref_pb > 0 and pb == 0: diffs.append("范文有换页,我没有")
        if ref_toc and not toc: diffs.append("范文有目录,我没有")
        if ref_head > 0 and head_styles == 0: diffs.append("范文有标题样式,我全Normal")
        R.append((len(diffs) == 0, "对范文/磋商结构指纹", ("对齐" if not diffs else "; ".join(diffs))))

    return R

def render_pages(docx_path, out_dir, dpi=110):
    """Gate B 第一步: 渲染成逐页 PNG, 返回路径列表(供 Claude 逐张 Read 去看)。"""
    os.makedirs(out_dir, exist_ok=True)
    for f in glob.glob(os.path.join(out_dir, "*")):
        try: os.remove(f)
        except: pass
    subprocess.run(["soffice", "--headless", "--convert-to", "pdf", "--outdir", out_dir, docx_path],
                   check=False, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL, timeout=180)
    pdf = glob.glob(os.path.join(out_dir, "*.pdf"))
    if not pdf: return []
    subprocess.run(["pdftoppm", "-png", "-r", str(dpi), pdf[0], os.path.join(out_dir, "page")],
                   check=False, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL, timeout=180)
    return sorted(glob.glob(os.path.join(out_dir, "page*.png")))

def run_gate(docx_path, reference=None, out_png_dir=None):
    print("=" * 74)
    print("Gate A · 结构闸门:", os.path.basename(docx_path))
    print("=" * 74)
    R = gate_a(docx_path, reference)
    passed = True
    for ok, name, detail in R:
        mark = "✅" if ok is True else ("❌" if ok is False else "👁")
        if ok is False: passed = False
        print(f" {mark} {name:<26} {detail}")
    print(f"\nGate A: {'全绿 ✅' if passed else '有红 ❌ —— 必须修到全绿'}   (👁 = 交给 Gate B 亲眼看)")
    if out_png_dir:
        pngs = render_pages(docx_path, out_png_dir)
        print(f"\nGate B · 渲染 {len(pngs)} 页 -> Claude 逐张 Read 去看(豆腐块/表格出界/标题孤行/图压字/页码):")
        for p in pngs: print("   ", p)
    return passed, R

if __name__ == "__main__":
    import sys
    ref = sys.argv[2] if len(sys.argv) > 2 else None
    out = sys.argv[3] if len(sys.argv) > 3 else "/tmp/gateB"
    run_gate(sys.argv[1], ref, out)

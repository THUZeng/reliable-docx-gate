# -*- coding: utf-8 -*-
"""
变异测试闸门本身 (mutation-test the gate) —— 信一个闸门前，先证明它能抓到已知缺陷。
往一份「已过闸门的好稿」里逐类植入 XXX占位 / 删换页 / 全改Normal / 偏离改不满足 / 删目录 / 孤儿样式 …
跑 Gate A，断言对应那条检查**变红**。哪类没红 = 闸门瞎了 → 先修闸门再验稿。

这就是当初漏 XXX 的解药：若早跑这个，"注入XXX→占位检查该红"一试就知道闸门根本不查占位。
研究依据: mutation testing (arxiv 2102.11378) —— 覆盖率保证不了什么，能抓到植入缺陷才算数。

用法:  python3 gate_selftest.py 已过闸门的好稿.docx
返回码 0 = 所有缺陷类都被抓到; 非0 = 有闸门盲区(打印哪类漏了)。
"""
import sys, os, copy, tempfile
from docx import Document
from docx.oxml.ns import qn
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from format_gate import gate_a

# ---- 每个变异子: 拿一个 Document, 植入一类缺陷 (原地改) ----
def m_inject_xxx(d):
    if d.tables: d.tables[0].rows[-1].cells[0].paragraphs[0].add_run("XXX")
    else: d.add_paragraph("XXX")
def m_inject_seepage(d):
    tgt = d.tables[0].rows[-1].cells[-1] if d.tables else None
    (tgt.paragraphs[0].add_run("见投标文件第页") if tgt else d.add_paragraph("见投标文件第页"))
def m_inject_emptyparen(d):
    d.add_paragraph("（此处填写项目负责人）")
def m_strip_pagebreaks(d):
    # 真删(不是置 val=0), 模拟「整篇不分页」的真实坏稿
    for pbb in d.element.body.findall('.//'+qn('w:pageBreakBefore')):
        pbb.getparent().remove(pbb)
    for br in d.element.body.findall('.//'+qn('w:br')):
        if br.get(qn('w:type')) == 'page': br.getparent().remove(br)
def m_flatten_styles(d):
    for p in d.paragraphs:
        try: p.style = d.styles['Normal']
        except Exception: pass
def m_break_deviation(d):
    for tb in d.tables:
        hdr = " ".join(c.text for c in tb.rows[0].cells)
        if "偏离" in hdr:
            col = next((j for j,c in enumerate(tb.rows[0].cells) if "偏离" in c.text), None)
            if col is not None and len(tb.rows) > 1:
                cell = tb.rows[1].cells[col]
                for r in cell.paragraphs[0].runs: r.text = ""
                cell.paragraphs[0].add_run("不满足")
            break
def m_strip_toc(d):
    for p in list(d.paragraphs):
        if p.text.strip() in ("目录","目 录"): p._p.getparent().remove(p._p); break
def m_orphan_style(d):
    st = d.styles.add_style('_orphan_test_', 1) if '_orphan_test_' not in [s.name for s in d.styles] else d.styles['_orphan_test_']
    if d.paragraphs: d.paragraphs[-1].style = st

# (变异名, 变异子, 期望变红的检查名子串)
MUTATIONS = [
    ("注入 XXX 占位",        m_inject_xxx,        "无残留占位符"),
    ("注入 见投标文件第页",   m_inject_seepage,    "无残留占位符"),
    ("注入 （此处填写…）",    m_inject_emptyparen, "无残留占位符"),
    ("删除所有换页",         m_strip_pagebreaks,  "有换页"),
    ("全部段落改 Normal",    m_flatten_styles,    "标题样式"),
    ("偏离表改「不满足」",    m_break_deviation,   "偏离表全满足"),
    ("删除目录",             m_strip_toc,         "有目录"),
    ("注入孤儿样式",         m_orphan_style,      "孤儿"),
]

def find(R, sub):
    return next((it for it in R if sub in it[1]), None)

def run(good_path):
    print("="*70); print("变异测试闸门:", os.path.basename(good_path)); print("="*70)
    # 先确认好稿本身是干净的(基线)
    base = gate_a(good_path)
    base_reds = [n for ok,n,_ in base if ok is False]
    print("基线(好稿) Gate A 红项:", base_reds or "无(干净)")
    blind = []
    for name, mutate, expect in MUTATIONS:
        base_item = find(base, expect)
        # 该检查在这份稿里根本不适用(如资格文件无偏离表 => 基线 None) => N/A, 不算盲区
        if base_item is not None and base_item[0] is None:
            print(f"  ➖ N/A     {name:<16} → 「{expect}」本稿不适用(需拿含该项的稿测)"); continue
        d = Document(good_path)
        try: mutate(d)
        except Exception as e:
            print(f"  ⚠ 变异[{name}]施加失败: {e!r} —— 跳过"); continue
        tmp = os.path.join(tempfile.gettempdir(), "gate_mut.docx"); d.save(tmp)
        R = gate_a(tmp)
        item = find(R, expect)
        caught = item is not None and item[0] is False
        was_red_at_base = base_item is not None and base_item[0] is False
        ok = caught and not was_red_at_base
        print(f"  {'✅ 抓到' if ok else '❌ 漏！闸门盲区'}  {name:<16} → 检查「{expect}」{'变红' if caught else '仍绿'}")
        if not ok: blind.append((name, expect))
    print("-"*70)
    if blind:
        print("闸门盲区(必须先补 format_gate 再验稿):")
        for n,e in blind: print(f"   ✗ {n} —— 「{e}」没抓到")
        return 1
    print("✅ 所有已知缺陷类都被 Gate A 抓到 —— 闸门本身合格，可用于验稿。")
    return 0

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("用法: python3 gate_selftest.py 已过闸门的好稿.docx"); sys.exit(2)
    sys.exit(run(sys.argv[1]))
